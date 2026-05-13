"""Egyszeri segédprogram: .docx fájlok szövegét kiveszi markdown formátumra."""
import sys
from docx import Document


def docx_to_markdown(path: str) -> str:
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name if para.style else "Normal"
        if "Heading 1" in style:
            lines.append(f"# {text}")
        elif "Heading 2" in style:
            lines.append(f"## {text}")
        elif "Heading 3" in style:
            lines.append(f"### {text}")
        elif "Heading 4" in style:
            lines.append(f"#### {text}")
        elif "Heading 5" in style:
            lines.append(f"##### {text}")
        elif "List Bullet" in style:
            lines.append(f"- {text}")
        elif "List Number" in style:
            lines.append(f"1. {text}")
        else:
            lines.append(text)
    for tbl_idx, table in enumerate(doc.tables, 1):
        lines.append(f"\n### Táblázat #{tbl_idx}\n")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("Használat: python extract-docx-text.py <docx-fájl> [kimenet.md]")
    src = sys.argv[1]
    md = docx_to_markdown(src)
    if len(sys.argv) >= 3:
        with open(sys.argv[2], "w", encoding="utf-8") as f:
            f.write(md)
        print(f"Kiírva: {sys.argv[2]} ({len(md)} karakter)")
    else:
        sys.stdout.buffer.write(md.encode("utf-8"))
