"""
Word dokumentumok batch feldolgozása
Beolvassa az összes .docx fájlt és kimenti .txt formátumba
"""
import os
import sys
from pathlib import Path

try:
    from docx import Document
    print("✓ python-docx imported successfully")
except ImportError:
    print("✗ python-docx not found. Installing...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document

def extract_text_from_docx(docx_path):
    """Word dokumentum szövegének kinyerése"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Táblázatok kezelése
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"HIBA: {str(e)}"

def process_all_docx_files(base_dir):
    """Összes .docx fájl feldolgozása"""
    base_path = Path(base_dir)
    output_dir = base_path / "extracted_texts"
    output_dir.mkdir(exist_ok=True)
    
    docx_files = list(base_path.rglob("*.docx"))
    
    print(f"\n📄 Talált {len(docx_files)} Word dokumentum")
    print(f"📁 Kimenet: {output_dir}\n")
    
    results = []
    
    for idx, docx_file in enumerate(docx_files, 1):
        relative_path = docx_file.relative_to(base_path)
        print(f"[{idx}/{len(docx_files)}] Feldolgozás: {relative_path}")
        
        # Szöveg kinyerése
        text_content = extract_text_from_docx(docx_file)
        
        # Kimenet fájlnév generálása
        output_filename = str(relative_path).replace(os.sep, "_").replace(".docx", ".txt")
        output_path = output_dir / output_filename
        
        # Mentés UTF-8 kódolással
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"=== {relative_path} ===\n\n")
                f.write(text_content)
            
            results.append({
                'file': str(relative_path),
                'output': str(output_path),
                'chars': len(text_content),
                'status': '✓'
            })
            print(f"  ✓ Mentve: {output_filename} ({len(text_content)} karakter)")
        except Exception as e:
            results.append({
                'file': str(relative_path),
                'output': None,
                'chars': 0,
                'status': f'✗ {str(e)}'
            })
            print(f"  ✗ Hiba: {e}")
    
    # Összesítő riport
    print("\n" + "="*60)
    print("FELDOLGOZÁS KÉSZ")
    print("="*60)
    
    success_count = sum(1 for r in results if r['status'] == '✓')
    total_chars = sum(r['chars'] for r in results)
    
    print(f"✓ Sikeres: {success_count}/{len(results)}")
    print(f"📝 Összesen: {total_chars:,} karakter")
    print(f"📁 Kimenet mappa: {output_dir}")
    
    # Részletes lista
    print("\nFeldolgozott fájlok:")
    for r in results:
        print(f"  {r['status']} {r['file']}")
    
    # Összevont fájl készítése
    print("\n📦 Összevont fájl készítése...")
    combined_path = output_dir / "_OSSZES_DOKUMENTUM_EGYUTT.txt"
    
    try:
        with open(combined_path, 'w', encoding='utf-8') as combined_file:
            combined_file.write("="*80 + "\n")
            combined_file.write("VALUTAVÁLTÓ PROGRAM - ÖSSZES DOKUMENTUM\n")
            combined_file.write(f"Generálva: {base_path}\n")
            combined_file.write("="*80 + "\n\n")
            
            for result in results:
                if result['status'] == '✓' and result['output']:
                    combined_file.write("\n" + "="*80 + "\n")
                    combined_file.write(f"FÁJL: {result['file']}\n")
                    combined_file.write("="*80 + "\n\n")
                    
                    with open(result['output'], 'r', encoding='utf-8') as f:
                        # Kihagyjuk a fejlécet az egyedi fájlokból
                        content = f.read()
                        # Töröljük az első 3 sort (=== header ===)
                        lines = content.split('\n')
                        combined_file.write('\n'.join(lines[3:]))
                    
                    combined_file.write("\n\n")
        
        print(f"✓ Összevont fájl létrehozva: {combined_path}")
        print(f"  Méret: {os.path.getsize(combined_path):,} byte")
    except Exception as e:
        print(f"✗ Hiba az összevont fájl készítésekor: {e}")
    
    return results

if __name__ == "__main__":
    # Beállítások
    base_directory = r"c:\Users\Kósa Zoltán\OneDrive\Dokumentumok\GIThub munkamappa\valutavalto-program\extracted"
    
    print("="*60)
    print("WORD DOKUMENTUMOK BATCH FELDOLGOZÓ")
    print("="*60)
    print(f"Mappa: {base_directory}\n")
    
    if not os.path.exists(base_directory):
        print(f"✗ HIBA: A mappa nem létezik: {base_directory}")
        sys.exit(1)
    
    results = process_all_docx_files(base_directory)
    
    print("\n✓ BEFEJEZVE")
